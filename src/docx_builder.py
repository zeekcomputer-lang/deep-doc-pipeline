"""
DOCX Builder — v3.1.

Builds a professional business report DOCX:
  Cover (1 page) + Body sections (flowing) + Implications (시사점)

v3.1 changes:
  - Body sections flow continuously (no forced page break between them)
  - Implications section appended after body with subtle divider
  - Korean title/subtitle support on cover page
  - Section dividers between body sections

Uses python-docx with lxml helpers for section properties.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ════════════════════════════════════════════════════════════════
# Color / Style constants
# ════════════════════════════════════════════════════════════════
NAVY = RGBColor(0x1B, 0x36, 0x5D)
GRAY_SUBTITLE = RGBColor(0x5B, 0x7B, 0x9A)
BODY_TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)
ACCENT_LINE_COLOR = "1B365D"
DIVIDER_COLOR = "CCCCCC"

FONT_PRIMARY = "맑은 고딕"
FONT_FALLBACK = "Calibri"


class DocxBuilder:
    """Build a professional DOCX: cover + body sections + implications."""

    def __init__(
        self,
        blueprint: dict,
        body_sections: list[str],
        implications_text: str,
        meta: dict,
    ):
        """
        Args:
            blueprint: DocumentBlueprint as dict.
            body_sections: [section_1_korean_md, section_2_korean_md]
            implications_text: Korean implications markdown (may start with ## 시사점)
            meta: from prompt_config.get_docx_meta() + kr title overrides
        """
        self.blueprint = blueprint
        self.body_sections = body_sections
        self.implications_text = implications_text
        self.meta = meta
        self.doc = Document()

    # ────────────────────────────────────────────────────────────
    # Public API
    # ────────────────────────────────────────────────────────────

    def build(self, output_path: str) -> str:
        """Build and save DOCX. Returns output path."""
        self._setup_default_styles()
        self._set_margins(self.doc.sections[0])

        # Cover page
        self._build_cover()

        # Body content — new page, sections flow continuously
        self._add_page_break_section()

        for i, section_text in enumerate(self.body_sections):
            if i > 0:
                self._add_section_divider()
            self._render_markdown(section_text)

        # Implications section
        if self.implications_text:
            self._add_section_divider()
            self._render_markdown(self.implications_text)

        # Headers & footers (not on cover)
        self._setup_headers_footers()

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        return output_path

    # ────────────────────────────────────────────────────────────
    # Style setup
    # ────────────────────────────────────────────────────────────

    def _setup_default_styles(self):
        """Configure default document styles."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = FONT_PRIMARY
        font.size = Pt(10.5)
        font.color.rgb = BODY_TEXT_COLOR
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} '
                               f'w:eastAsia="{FONT_PRIMARY}" '
                               f'w:hAnsi="{FONT_FALLBACK}"/>')
            rpr.append(rfonts)
        else:
            rfonts.set(qn("w:eastAsia"), FONT_PRIMARY)
            rfonts.set(qn("w:hAnsi"), FONT_FALLBACK)

    def _set_margins(self, section):
        """Set 2.54cm margins on all sides."""
        margin = Cm(2.54)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # ────────────────────────────────────────────────────────────
    # Cover page
    # ────────────────────────────────────────────────────────────

    def _build_cover(self):
        """Build a centered cover page with title, subtitle, accent line, date, org."""
        doc_title = self.meta.get("title") or self.blueprint.get("doc_title", "Untitled")
        doc_subtitle = self.meta.get("subtitle") or self.blueprint.get("doc_subtitle", "")
        organization = self.meta.get("organization", "")

        # Spacer to push title toward center
        for _ in range(6):
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = 1.0

        # Title
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(12)
        run = p_title.add_run(doc_title)
        run.font.size = Pt(28)
        run.font.color.rgb = NAVY
        run.font.name = FONT_PRIMARY
        run.bold = True
        self._set_run_east_asian_font(run)

        # Accent line
        self._add_accent_line()

        # Subtitle
        if doc_subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = Pt(8)
            p_sub.paragraph_format.space_after = Pt(4)
            run = p_sub.add_run(doc_subtitle)
            run.font.size = Pt(14)
            run.font.color.rgb = GRAY_SUBTITLE
            run.font.name = FONT_PRIMARY
            self._set_run_east_asian_font(run)

        # Push date/org to bottom
        for _ in range(8):
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = 1.0

        # Date
        today_str = datetime.now().strftime("%Y년 %m월 %d일")
        p_date = self.doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.paragraph_format.space_after = Pt(4)
        run = p_date.add_run(today_str)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_SUBTITLE
        run.font.name = FONT_PRIMARY
        self._set_run_east_asian_font(run)

        # Organization
        if organization:
            p_org = self.doc.add_paragraph()
            p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_org.paragraph_format.space_after = Pt(0)
            run = p_org.add_run(organization)
            run.font.size = Pt(11)
            run.font.color.rgb = GRAY_SUBTITLE
            run.font.name = FONT_PRIMARY
            self._set_run_east_asian_font(run)

    def _add_accent_line(self):
        """Add a centered accent line (navy bottom border)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="1" '
            f'           w:color="{ACCENT_LINE_COLOR}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        ind = parse_xml(
            f'<w:ind {nsdecls("w")} w:left="2880" w:right="2880"/>'
        )
        pPr.append(ind)

    # ────────────────────────────────────────────────────────────
    # Section dividers
    # ────────────────────────────────────────────────────────────

    def _add_section_divider(self):
        """Add a subtle thin gray line between content sections."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(16)

        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" '
            f'           w:color="{DIVIDER_COLOR}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ────────────────────────────────────────────────────────────
    # Markdown rendering (body pages)
    # ────────────────────────────────────────────────────────────

    def _render_markdown(self, md_text: str):
        """Parse markdown text and render into DOCX paragraphs.

        Supports: ## headings, ### subheadings, - bullet lists, **bold**, `code`.
        """
        if not md_text:
            return

        lines = md_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Heading (## or ###)
            heading_match = re.match(r'^(#{2,3})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14) if level == 2 else Pt(8)
                p.paragraph_format.space_after = Pt(8) if level == 2 else Pt(6)
                run = p.add_run(text)
                run.font.size = Pt(16) if level == 2 else Pt(13)
                run.font.color.rgb = NAVY
                run.bold = True
                run.font.name = FONT_PRIMARY
                self._set_run_east_asian_font(run)
                i += 1
                continue

            # Bullet list
            bullet_match = re.match(r'^[-*]\s+(.+)$', stripped)
            if bullet_match:
                text = bullet_match.group(1)
                p = self.doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                self._add_inline_formatted_runs(p, text)
                i += 1
                continue

            # Regular paragraph — collect continuation lines
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    i += 1
                    break
                if re.match(r'^(#{2,3})\s+', next_line):
                    break
                if re.match(r'^[-*]\s+', next_line):
                    break
                para_lines.append(next_line)
                i += 1

            full_text = " ".join(para_lines)
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            self._add_inline_formatted_runs(p, full_text)

    def _add_inline_formatted_runs(self, paragraph, text: str):
        """Parse inline markdown (**bold**, `code`) and add as runs."""
        pattern = re.compile(r'(\*\*.*?\*\*|`[^`]+`)')
        parts = pattern.split(text)

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                inner = part[2:-2]
                run = paragraph.add_run(inner)
                run.bold = True
                run.font.name = FONT_PRIMARY
                run.font.size = Pt(10.5)
                run.font.color.rgb = BODY_TEXT_COLOR
                self._set_run_east_asian_font(run)
            elif part.startswith("`") and part.endswith("`"):
                inner = part[1:-1]
                run = paragraph.add_run(inner)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                run = paragraph.add_run(part)
                run.font.name = FONT_PRIMARY
                run.font.size = Pt(10.5)
                run.font.color.rgb = BODY_TEXT_COLOR
                self._set_run_east_asian_font(run)

    # ────────────────────────────────────────────────────────────
    # Section / Page breaks
    # ────────────────────────────────────────────────────────────

    def _add_page_break_section(self):
        """Add a new section with page break (NEW_PAGE start type)."""
        new_section = self.doc.add_section(WD_SECTION_START.NEW_PAGE)
        self._set_margins(new_section)
        return new_section

    # ────────────────────────────────────────────────────────────
    # Headers & Footers
    # ────────────────────────────────────────────────────────────

    def _setup_headers_footers(self):
        """Set up headers and footers.

        Cover page (section 0): no header/footer (different first page).
        Body pages (sections 1+): header with doc title, footer with page number.
        """
        doc_title = self.meta.get("title") or self.blueprint.get("doc_title", "Untitled")

        # Section 0 (cover): enable "different first page" for blank header/footer
        section_0 = self.doc.sections[0]
        sectPr = section_0._sectPr
        titlePg = sectPr.find(qn("w:titlePg"))
        if titlePg is None:
            titlePg = parse_xml(f'<w:titlePg {nsdecls("w")}/>')
            sectPr.append(titlePg)

        # Body sections: header + footer
        for sec_idx in range(1, len(self.doc.sections)):
            section = self.doc.sections[sec_idx]
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            # Header: document title (right-aligned, small)
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.text = ""
            run = hp.add_run(doc_title)
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY_SUBTITLE
            run.font.name = FONT_PRIMARY
            self._set_run_east_asian_font(run)
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Footer: page number (centered)
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = ""
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_page_number_field(fp)

    def _add_page_number_field(self, paragraph):
        """Add a PAGE field to a paragraph for automatic page numbering."""
        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_SUBTITLE

        fldChar_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        run._element.append(fldChar_begin)

        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run._element.append(instrText)

        fldChar_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run._element.append(fldChar_end)

    # ────────────────────────────────────────────────────────────
    # Helpers
    # ────────────────────────────────────────────────────────────

    @staticmethod
    def _set_run_east_asian_font(run):
        """Set East Asian font on a run element for Korean text rendering."""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_PRIMARY}"/>'
            )
            rPr.append(rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), FONT_PRIMARY)
